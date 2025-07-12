import re
import tempfile
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt

app = FastAPI()

class Message(BaseModel):
    speaker: str
    text: str
    timestamp: str | None = None

class CompileRequest(BaseModel):
    title: str
    messages: list[Message]

def slugify(text: str) -> str:
    return re.sub(r'[^A-Za-z0-9]+', '_', text).strip('_') or 'document'

def add_html_to_doc(doc: Document, html: str):
    """
    Конвертує HTML (згенерований markdown) у абзаци та стилі DOCX.
    """
    soup = BeautifulSoup(html, "html.parser")
    for elem in soup.children:
        if isinstance(elem, NavigableString):
            doc.add_paragraph(str(elem))
        elif not isinstance(elem, Tag):
            continue
        # заголовки
        if elem.name == "h1":
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name in ("h2", "h3"):
            level = 3 if elem.name == "h3" else 3
            doc.add_heading(elem.get_text(), level=level)
        # списки
        elif elem.name == "ul":
            for li in elem.find_all("li", recursive=False):
                p = doc.add_paragraph(li.get_text(), style="List Bullet")
        elif elem.name == "ol":
            for li in elem.find_all("li", recursive=False):
                p = doc.add_paragraph(li.get_text(), style="List Number")
        # блок коду
        elif elem.name == "pre":
            code = elem.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        # параграф із внутрішнім форматуванням
        else:
            p = doc.add_paragraph()
            def recurse(node):
                if isinstance(node, NavigableString):
                    p.add_run(str(node))
                elif isinstance(node, Tag):
                    text = node.get_text()
                    run = p.add_run(text)
                    if node.name in ("strong", "b"):
                        run.bold = True
                    if node.name in ("em", "i"):
                        run.italic = True
                    if node.name == "code":
                        run.font.name = "Courier New"
                        run.font.size = Pt(10)
                    # рекурсивно для вкладених елементів
                    for child in node.contents:
                        recurse(child)
            for child in elem.contents:
                recurse(child)

@app.post("/compileChatToDocx")
def compile_chat(req: CompileRequest):
    """
    Приймає title та масив повідомлень {speaker, text, timestamp},
    формує DOCX з розділом на кожне повідомлення, зберігає стилі Markdown та емодзі,
    повертає файл через FileResponse.
    """
    try:
        doc = Document()
        # Титулка
        doc.add_heading(req.title, level=1)

        for msg in req.messages:
            # header: Speaker [timestamp]:
            header = msg.speaker
            if msg.timestamp:
                header += f" [{msg.timestamp}]"
            p = doc.add_paragraph()
            r = p.add_run(header + ": ")
            r.bold = True

            # конвертуємо Markdown → HTML → DOCX
            html = markdown.markdown(
                msg.text, 
                extensions=["extra", "sane_lists", "fenced_code"]
            )
            # додаємо внутрішню частину тієї ж абзаци
            add_html_to_doc(doc, html)

        # зберігаємо тимчасово
        slug = slugify(req.title)
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        path = tmp.name
        tmp.close()
        doc.save(path)

        return FileResponse(
            path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{slug}.docx"
        )
    except Exception as e:
        raise HTTPException(500, detail=str(e))
